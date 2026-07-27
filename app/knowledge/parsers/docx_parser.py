"""
app/knowledge/parsers/docx_parser.py
──────────────────────────────────────
DOCX text extraction using python-docx.
Extracts all paragraph text; tables are joined with spaces.
"""

import logging
from pathlib import Path

from .base_parser import BaseParser, ParsedPage

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):

    def parse(self, file_path: str) -> list[ParsedPage]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required: pip install python-docx")

        path = Path(file_path)
        if not path.exists():
            raise ValueError(f"File not found: {file_path}")

        try:
            doc   = Document(str(path))
            parts: list[str] = []

            # Paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            # Tables — each row flattened to a space-joined string
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)

            full_text = "\n".join(parts)
            if not full_text.strip():
                logger.warning("DOCX produced no text: %s", path.name)
                return []

            return [ParsedPage(text=full_text, page_number=0)]

        except Exception as exc:
            logger.error("DOCX parse error: %s | %s", path.name, exc)
            raise ValueError(f"Failed to parse DOCX '{path.name}': {exc}") from exc
